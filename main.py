# -*- coding: utf-8 -*-
"""
Created on Thu May  6 19:57:23 2021

@author: Owner
"""

# -*- coding: utf-8 -*-
"""
Created on Tue Apr  6 00:06:22 2021

@author: Owner
"""

# Import docx NOT python-docx
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
import copy
from docx2pdf import convert

#%% Function Zone

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def replaceParaph(doc,search="",replace=""):
    for paragraph in doc.paragraphs:
        if search in paragraph.text:
            print(paragraph.text)
            #paragraph.text = 'new text containing ocean'
            paragraph.text = str(paragraph.text).replace(search,replace)

def add_keyword(keyword):
    keyword = str(keyword)
    keyword = "{{" + keyword + "}}"
    return keyword

def create_replacer(data:pd.DataFrame,row:int):
    replacer_ = data.loc[row].to_dict()
    
    for old in replacer_.keys():
        
        replacer_[add_keyword(old)] = replacer_.pop(old)
        
    return replacer_

def keyword_replacer_document(doc :docx.Document, replacer_dict:dict):
    new_doc = copy.deepcopy(doc)
    for paragraph in new_doc.paragraphs:
        for keyword in replacer_dict.keys():
            if keyword in paragraph.text:
                #paragraph.text = 'new text containing ocean'
                paragraph.text = str(paragraph.text).replace(keyword,replacer_dict[keyword])
            else: pass
    print(replacer_dict)
    return new_doc




    
def multi_report(data:pd.DataFrame,main_name:str,path=""): # ,doc:docx.document.Document
    
    nro_report = data.shape[0]
    
    if main_name == "":main_name = "document"
    for nro in range(0,nro_report):
        print(create_replacer(data,nro))
        # print(replacer_dict)
        doc = docx.Document("in_word/Informe Psicolaboral_2020.docx")
        new_doc = keyword_replacer_document(doc ,create_replacer(data,row = nro))
        new_doc.save('output/{name}_{nro}.docx'.format(name=main_name,nro=nro)) 
        # keyword_replacer_document(doc ,create_replacer(data,row = nro)).save('output/{name}_{nro}.docx'.format(name=main_name,nro=nro))

def multi_report_n(data:pd.DataFrame,doc:docx.document.Document,main_name:str): # 
    
    nro_report = data.shape[0]
    # new_doc = copy.copy(doc)
    if main_name == "":main_name = "document"
    for nro in range(0,nro_report):
        print(create_replacer(data,nro))
        # print(replacer_dict)
        # doc = docx.Document("in_word/Informe Psicolaboral_2020.docx")
        new_doc = keyword_replacer_document(doc,create_replacer(data,row = nro))
        new_doc.save('output/{name}_{nro}.docx'.format(name=main_name,nro=nro))


# convert("input.docx")
# convert("input.docx", "output.pdf")
# convert("my_docx_folder/")
#%%
# styles = document.styles

## Styles
# Paragraph
document = Document()
style = document.styles
style = style.add_style('Global_text', WD_STYLE_TYPE.PARAGRAPH)
font = style.font
font.name = 'Nunito'
font.size = Pt(12)
font.bold = True

# Create an instance of a word document
doc = docx.Document("in_word/Informe Psicolaboral_2020.docx")

# text = getText("in_word/Informe Psicolaboral_2020.docx")
data = pd.read_excel("in_excel/clientes.xlsx",sheet_name=0)

#%% 

doc = docx.Document("in_word/Informe Psicolaboral_2020.docx")
multi_report_n(data,doc,"reporte")

#%%

multi_report(data,"reporte",path = "in_word/Informe Psicolaboral_2020.docx" )
